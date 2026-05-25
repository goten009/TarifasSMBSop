from pathlib import Path
from io import BytesIO
from datetime import datetime
import re
import unicodedata
import zipfile
import copy

import pandas as pd
import streamlit as st

from PIL import Image
from docx import Document


# ============================================================
# CONFIGURACIÓN GENERAL
# ============================================================

st.set_page_config(
    page_title="Cotizador Sophos SMB | Media Commerce",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="collapsed",
)

BASE_DIR = Path(__file__).resolve().parent

PRICE_FILE = BASE_DIR / "data" / "Lista de precios SOPHOS SMB.xlsx"
LOGO_FILE = BASE_DIR / "assets" / "logo_mc.png"

TEMPLATE_OFERTA_UTM_FILE = BASE_DIR / "assets" / "oferta_venta_utm.docx"
TEMPLATE_SERVICIO_NGFW_FILE = BASE_DIR / "assets" / "oferta_servicio_ngfw.docx"
TEMPLATE_ENDPOINT_FILE = BASE_DIR / "assets" / "oferta_endpoint.docx"

WORD_MEDIA_FIREWALL_IMAGE_VENTA = "word/media/image2.png"
WORD_MEDIA_FIREWALL_IMAGE_SERVICIO = "word/media/image2.jpeg"

FIREWALL_IMAGE_FILES = {
    "XGS 88": BASE_DIR / "assets" / "XGS88.jpg",
    "XGS 108": BASE_DIR / "assets" / "XGS108.jpg",
    "XGS 118": BASE_DIR / "assets" / "XGS118.jpg",
    "XGS 128": BASE_DIR / "assets" / "XGS128.jpg",
    "XGS 138": BASE_DIR / "assets" / "XGS138.jpg",
}

FAMILIAS_VALIDAS = [
    "Firewall",
    "Endpoint",
    "Server",
    "EDR User",
    "EDR Server",
    "XDR User",
    "XDR Server",
]

VIGENCIAS_VALIDAS = [12, 24, 36]

TIPOS_VENTA_FIREWALL = [
    "Venta one shot",
    "Servicio mensual",
]

TIPOS_INSTALACION_LICENCIAS = [
    "Sin instalación",
    "Instalación básica (5 licencias)",
    "Instalación completa",
]

TIPOS_INSTALACION_FIREWALL_SERVICIO = [
    "Sin cobro de instalación",
    "Con cobro de instalación",
]

TIPOS_INSTALACION_FIREWALL_VENTA = [
    "Sin instalación",
    "Instalación básica",
    "Instalación completa",
]


# ============================================================
# VARIABLES FINANCIERAS
# ============================================================

IMPUESTOS = 0.012
GASTOS = 0.045
CARGA_PRESTACIONAL_COMISION = 0.014
COMISION = 0.035

PORCENTAJE_COSTO_SOBRE_VENTA = 0.739
PORCENTAJE_COSTO_INSTALACION = 0.10


# ============================================================
# TARIFAS FIREWALL SERVICIO
# ============================================================

TARIFA_SERVICIO_FIREWALL = {
    "XGS 88": {12: 367000, 24: 314000, 36: 294000},
    "XGS 108": {12: 447000, 24: 382000, 36: 357000},
    "XGS 118": {12: 563000, 24: 493000, 36: 449000},
    "XGS 128": {12: 758000, 24: 651000, 36: 595000},
    "XGS 138": {12: 927000, 24: 846000, 36: 785000},
}

NRC_SERVICIO_FIREWALL = {
    "Sin cobro de instalación": 0,
    "Con cobro de instalación": 350000,
}

CAPACIDAD_FIREWALL = {
    "XGS 88": 20,
    "XGS 108": 40,
    "XGS 118": 70,
    "XGS 128": 100,
    "XGS 138": 200,
}

MODELO_TABLA_SERVICIO = {
    "XGS 88": "XGS88",
    "XGS 108": "XGS108",
    "XGS 118": "XGS118",
    "XGS 128": "XGS128",
    "XGS 138": "XGS138",
}


# ============================================================
# ESTILO CORPORATIVO
# ============================================================

st.markdown(
    """
    <style>
        :root {
            --mc-blue: #0057A8;
            --mc-blue-dark: #003B73;
            --mc-blue-soft: #E8F2FC;
            --mc-yellow: #FFC400;
            --mc-gray-bg: #F5F7FA;
            --mc-gray-line: #D9E1EA;
            --mc-text: #1F2937;
            --mc-muted: #6B7280;
            --mc-white: #FFFFFF;
            --mc-green: #2E7D32;
        }

        .stApp {
            background: linear-gradient(180deg, #F7FAFD 0%, #FFFFFF 38%);
            color: var(--mc-text);
        }

        [data-testid="stToolbar"] {
            visibility: hidden;
            height: 0%;
            position: fixed;
        }

        [data-testid="stDecoration"] {
            background: linear-gradient(90deg, #0057A8 0%, #006FCB 55%, #FFC400 100%);
            height: 5px;
        }

        .block-container {
            padding-top: 1.5rem;
            padding-bottom: 3rem;
            max-width: 1320px;
        }

        .mc-header {
            background: linear-gradient(135deg, #FFFFFF 0%, #EEF6FF 100%);
            border: 1px solid #D9E7F5;
            border-radius: 22px;
            padding: 28px 32px;
            margin-bottom: 24px;
            box-shadow: 0 12px 30px rgba(0, 57, 115, 0.08);
        }

        .mc-title {
            color: #003B73;
            font-size: 34px;
            line-height: 1.15;
            font-weight: 800;
            margin-bottom: 8px;
        }

        .mc-subtitle {
            color: #52616F;
            font-size: 16px;
            line-height: 1.45;
            max-width: 760px;
        }

        .mc-badge-row {
            display: flex;
            flex-wrap: wrap;
            gap: 8px;
            margin-top: 16px;
        }

        .mc-badge {
            background: #FFFFFF;
            color: #0057A8;
            border: 1px solid #BFD7EF;
            padding: 7px 12px;
            border-radius: 999px;
            font-size: 13px;
            font-weight: 700;
        }

        .mc-section-title {
            color: #003B73;
            font-size: 25px;
            font-weight: 800;
            margin: 18px 0 8px 0;
        }

        .mc-section-caption {
            color: #6B7280;
            font-size: 14px;
            margin-bottom: 18px;
        }

        .mc-panel {
            background: #FFFFFF;
            border: 1px solid #D9E1EA;
            border-radius: 20px;
            padding: 24px;
            margin-bottom: 24px;
            box-shadow: 0 10px 28px rgba(31, 41, 55, 0.06);
        }

        .mc-result-card {
            background: #FFFFFF;
            border: 1px solid #D9E1EA;
            border-left: 6px solid #0057A8;
            border-radius: 18px;
            padding: 18px 20px;
            min-height: 128px;
            box-shadow: 0 8px 24px rgba(0, 57, 115, 0.06);
        }

        .mc-result-card-yellow {
            border-left-color: #FFC400;
        }

        .mc-result-card-green {
            border-left-color: #2E7D32;
        }

        .mc-result-label {
            color: #6B7280;
            font-size: 14px;
            font-weight: 700;
            margin-bottom: 8px;
        }

        .mc-result-value {
            color: #1F2937;
            font-size: 30px;
            font-weight: 800;
            line-height: 1.2;
        }

        .mc-result-note {
            color: #6B7280;
            font-size: 12px;
            margin-top: 8px;
        }

        .mc-alert {
            background: #E8F2FC;
            border: 1px solid #BFD7EF;
            border-left: 6px solid #0057A8;
            color: #003B73;
            border-radius: 16px;
            padding: 16px 18px;
            font-weight: 650;
            margin: 14px 0 22px 0;
        }

        .mc-footer-note {
            color: #6B7280;
            font-size: 12px;
            text-align: center;
            margin-top: 28px;
            padding-top: 20px;
            border-top: 1px solid #D9E1EA;
        }

        div[data-testid="stDataFrame"] {
            border-radius: 16px;
            overflow: hidden;
        }

        .stSelectbox label, .stNumberInput label, .stCheckbox label {
            color: #374151 !important;
            font-weight: 700 !important;
        }
    </style>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# FUNCIONES UTILITARIAS
# ============================================================

def normalizar_texto(valor):
    if pd.isna(valor):
        return ""

    texto = str(valor).strip().lower()
    texto = unicodedata.normalize("NFKD", texto)
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = re.sub(r"[^a-z0-9]", "", texto)
    return texto


def convertir_numero(valor):
    if pd.isna(valor):
        return 0.0

    if isinstance(valor, (int, float)):
        return float(valor)

    texto = str(valor).strip()
    texto = texto.replace("USD", "")
    texto = texto.replace("$", "")
    texto = texto.replace(" ", "")

    if "," in texto and "." in texto:
        texto = texto.replace(".", "").replace(",", ".")
    elif "," in texto:
        texto = texto.replace(",", ".")

    try:
        return float(texto)
    except ValueError:
        return 0.0


def formato_cop(valor):
    return f"${valor:,.0f}".replace(",", ".")


def formato_usd(valor):
    return f"USD {valor:,.2f}"


def formato_porcentaje(valor):
    return f"{valor * 100:.1f}%".replace(".", ",")


def obtener_columna(df, posibles_nombres):
    columnas_normalizadas = {
        normalizar_texto(col): col
        for col in df.columns
    }

    for nombre in posibles_nombres:
        nombre_norm = normalizar_texto(nombre)
        if nombre_norm in columnas_normalizadas:
            return columnas_normalizadas[nombre_norm]

    return None


def sugerir_firewall(cantidad_usuarios):
    if cantidad_usuarios <= 20:
        return "XGS 88", 20
    elif cantidad_usuarios <= 40:
        return "XGS 108", 40
    elif cantidad_usuarios <= 70:
        return "XGS 118", 70
    elif cantidad_usuarios <= 100:
        return "XGS 128", 100
    elif cantidad_usuarios <= 200:
        return "XGS 138", 200
    return None, None


def calcular_precio_venta_desde_costo(costo_cop):
    if PORCENTAJE_COSTO_SOBRE_VENTA > 0:
        return costo_cop / PORCENTAJE_COSTO_SOBRE_VENTA
    return 0


def obtener_costo_usd_por_escenario(msrp_usd, dr_usd, escenario):
    msrp_usd = convertir_numero(msrp_usd)
    dr_usd = convertir_numero(dr_usd)

    if escenario == "Con DR aprobado" and dr_usd > 0:
        return dr_usd

    return msrp_usd


def render_result_card(label, value, note="", variant="blue"):
    css_class = "mc-result-card"
    if variant == "yellow":
        css_class += " mc-result-card-yellow"
    elif variant == "green":
        css_class += " mc-result-card-green"

    st.markdown(
        f"""
        <div class="{css_class}">
            <div class="mc-result-label">{label}</div>
            <div class="mc-result-value">{value}</div>
            <div class="mc-result-note">{note}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


# ============================================================
# FUNCIONES WORD GENERALES
# ============================================================

def obtener_imagen_firewall(modelo):
    return FIREWALL_IMAGE_FILES.get(modelo)


def convertir_imagen_a_bytes(image_path, formato="PNG"):
    with Image.open(image_path) as img:
        img = img.convert("RGB")
        buffer = BytesIO()

        if formato.upper() == "JPEG":
            img.save(buffer, format="JPEG", quality=95)
        else:
            img.save(buffer, format="PNG")

        buffer.seek(0)
        return buffer.read()


def aplicar_reemplazos_xml(texto_xml, reemplazos):
    texto_modificado = texto_xml

    for buscar, reemplazar in sorted(reemplazos.items(), key=lambda x: len(x[0]), reverse=True):
        texto_modificado = texto_modificado.replace(buscar, reemplazar)

    return texto_modificado


def reemplazar_imagen_y_xml_docx(docx_buffer, image_path, reemplazos_xml, media_target=None, image_format="PNG"):
    docx_buffer.seek(0)
    output = BytesIO()

    nueva_imagen = None

    if image_path is not None and Path(image_path).exists():
        nueva_imagen = convertir_imagen_a_bytes(image_path, formato=image_format)

    with zipfile.ZipFile(docx_buffer, "r") as zin:
        with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename.endswith(".xml") and item.filename.startswith("word/"):
                    try:
                        xml_text = data.decode("utf-8")
                        xml_text = aplicar_reemplazos_xml(xml_text, reemplazos_xml)
                        data = xml_text.encode("utf-8")
                    except UnicodeDecodeError:
                        pass

                if nueva_imagen is not None and media_target and item.filename == media_target:
                    data = nueva_imagen

                zout.writestr(item, data)

    output.seek(0)
    return output


def reemplazar_texto_parrafo(paragraph, reemplazos):
    texto_original = paragraph.text

    if not texto_original:
        return

    texto_nuevo = texto_original

    for buscar, reemplazar in sorted(reemplazos.items(), key=lambda x: len(x[0]), reverse=True):
        texto_nuevo = texto_nuevo.replace(buscar, reemplazar)

    if texto_nuevo == texto_original:
        return

    estilo = paragraph.style
    alineacion = paragraph.alignment

    for run in paragraph.runs:
        run.text = ""

    if paragraph.runs:
        paragraph.runs[0].text = texto_nuevo
    else:
        paragraph.add_run(texto_nuevo)

    paragraph.style = estilo
    paragraph.alignment = alineacion


def reemplazar_texto_documento(doc, reemplazos):
    for paragraph in doc.paragraphs:
        reemplazar_texto_parrafo(paragraph, reemplazos)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    reemplazar_texto_parrafo(paragraph, reemplazos)


def limpiar_celda(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.text = ""

    if not cell.paragraphs:
        cell.add_paragraph()


def escribir_celda(cell, texto):
    limpiar_celda(cell)
    paragraph = cell.paragraphs[0]

    if paragraph.runs:
        paragraph.runs[0].text = str(texto)
    else:
        paragraph.add_run(str(texto))


def asegurar_columnas_tabla(table, cantidad_columnas):
    """
    Mantiene compatibilidad si alguna plantilla antigua no tiene las columnas requeridas.
    Si la plantilla ya tiene 4 columnas, no hace nada.
    """
    while len(table.columns) < cantidad_columnas:
        ultima_col_idx = len(table.columns) - 1

        ancho_referencia = table.columns[ultima_col_idx].width
        if ancho_referencia is None:
            ancho_referencia = 1200000

        table.add_column(ancho_referencia)

        nueva_col_idx = len(table.columns) - 1

        for row in table.rows:
            celda_origen = row.cells[ultima_col_idx]
            celda_destino = row.cells[nueva_col_idx]

            tc_pr_origen = celda_origen._tc.get_or_add_tcPr()
            tc_pr_destino = celda_destino._tc.get_or_add_tcPr()

            for child in list(tc_pr_destino):
                tc_pr_destino.remove(child)

            for child in tc_pr_origen:
                tc_pr_destino.append(copy.deepcopy(child))

            if celda_origen.paragraphs and celda_destino.paragraphs:
                celda_destino.paragraphs[0].style = celda_origen.paragraphs[0].style
                celda_destino.paragraphs[0].alignment = celda_origen.paragraphs[0].alignment


def eliminar_filas_extra(table, filas_deseadas):
    while len(table.rows) > filas_deseadas:
        row = table.rows[-1]
        table._tbl.remove(row._tr)


def nombre_archivo_oferta(modelo, modalidad):
    modalidad_limpia = normalizar_texto(modalidad)
    fecha = datetime.now().strftime("%Y%m%d")
    modelo_limpio = str(modelo).replace(" ", "_")
    return f"Oferta_Sophos_{modelo_limpio}_{modalidad_limpia}_{fecha}.docx"


# ============================================================
# WORD - VENTA UTM FIREWALL
# ============================================================

def encontrar_tabla_plataforma(doc):
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 2:
            textos = []

            for row in table.rows:
                for cell in row.cells:
                    textos.append(cell.text.strip())

            joined = " ".join(textos).lower()

            if "plataforma" in joined and "licencia" in joined:
                return table

    return None


def encontrar_tabla_economica_venta(doc):
    for table in doc.tables:
        if len(table.rows) >= 2 and len(table.columns) >= 4:
            encabezados = [cell.text.strip().lower() for cell in table.rows[0].cells]

            if (
                "descripción" in encabezados[0]
                and "cant" in encabezados[1]
                and "valor unitario" in encabezados[2]
                and "valor total" in encabezados[3]
            ):
                return table

    return None


def actualizar_tabla_plataforma(doc, modelo, vigencia):
    table = encontrar_tabla_plataforma(doc)

    if table is None:
        return

    if len(table.rows) >= 1:
        escribir_celda(table.rows[0].cells[0], "Plataforma")
        escribir_celda(table.rows[0].cells[1], f"Sophos {modelo}")

    if len(table.rows) >= 2:
        escribir_celda(table.rows[1].cells[0], "Licenciamiento")
        escribir_celda(table.rows[1].cells[1], f"Licencia Xstream Protection {vigencia} meses")


def actualizar_tabla_economica_venta_utm(doc, datos):
    table = encontrar_tabla_economica_venta(doc)

    if table is None:
        return False

    filas = [
        [
            f"Sophos {datos['modelo']}",
            "1",
            formato_cop(datos["venta_equipo"]),
            formato_cop(datos["venta_equipo"]),
        ],
        [
            f"Licencia Xstream Protection {datos['vigencia']} meses",
            "1",
            formato_cop(datos["venta_licencia"]),
            formato_cop(datos["venta_licencia"]),
        ],
    ]

    if datos.get("venta_instalacion", 0) > 0:
        filas.append(
            [
                "Instalación y puesta en marcha",
                "1",
                formato_cop(datos["venta_instalacion"]),
                formato_cop(datos["venta_instalacion"]),
            ]
        )

    filas_necesarias = 1 + len(filas)

    while len(table.rows) < filas_necesarias:
        table.add_row()

    eliminar_filas_extra(table, filas_necesarias)

    for idx, fila in enumerate(filas, start=1):
        for col_idx, valor in enumerate(fila):
            escribir_celda(table.rows[idx].cells[col_idx], valor)

    return True


def generar_oferta_venta_utm_word(datos):
    if not TEMPLATE_OFERTA_UTM_FILE.exists():
        return None

    doc = Document(str(TEMPLATE_OFERTA_UTM_FILE))
    fecha_actual = datetime.now().strftime("%d/%m/%Y")

    reemplazos = {
        "El cliente INDEX TECNOLOGIA requiere": "El cliente requiere",
        "El cliente INDEX TECNOLOGÍA requiere": "El cliente requiere",
        "INDEX TECNOLOGIA": "Cliente",
        "INDEX TECNOLOGÍA": "Cliente",
        "Pereira, 21 de nov. de 25": f"Pereira, {fecha_actual}",
        "Edinson Rodriguez": "XXXX",
        "Edinson Rodríguez": "XXXX",
        "Edinson.rodriguez@mc.net.co": "XXXX",
        "Edinson.Rodriguez@mc.net.co": "XXXX",
        "edinson.rodriguez@mc.net.co": "XXXX",
        "Sophos XGS 108": f"Sophos {datos['modelo']}",
        "Sophos XGS 116": f"Sophos {datos['modelo']}",
        "Licencia Xtream protection 12 meses": f"Licencia Xstream Protection {datos['vigencia']} meses",
        "Licencia Xtream Protection": f"Licencia Xstream Protection {datos['vigencia']} meses",
        "Licencia Xstream Protection 12 meses": f"Licencia Xstream Protection {datos['vigencia']} meses",
    }

    reemplazar_texto_documento(doc, reemplazos)
    actualizar_tabla_plataforma(doc, datos["modelo"], datos["vigencia"])
    actualizar_tabla_economica_venta_utm(doc, datos)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    imagen_firewall = obtener_imagen_firewall(datos["modelo"])

    buffer = reemplazar_imagen_y_xml_docx(
        docx_buffer=buffer,
        image_path=imagen_firewall,
        reemplazos_xml=reemplazos,
        media_target=WORD_MEDIA_FIREWALL_IMAGE_VENTA,
        image_format="PNG",
    )

    buffer.seek(0)
    return buffer


# ============================================================
# WORD - SERVICIO NGFW FIREWALL
# ============================================================

def encontrar_tabla_preciario_servicio(doc):
    for table in doc.tables:
        texto = " ".join(cell.text.strip() for row in table.rows for cell in row.cells)
        texto_norm = normalizar_texto(texto)

        if "preciariongfwsophos" in texto_norm:
            return table

        if "mrr1ano" in texto_norm and "mrr2anos" in texto_norm and "mrr3anos" in texto_norm:
            return table

        if "xgs88" in texto_norm and "xgs108" in texto_norm and "xgs118" in texto_norm:
            return table

    return None


def actualizar_tabla_preciario_servicio(doc, datos):
    table = encontrar_tabla_preciario_servicio(doc)

    if table is None:
        return False

    modelo = datos["modelo"]
    modelo_tabla = MODELO_TABLA_SERVICIO.get(modelo, modelo.replace(" ", ""))
    capacidad = CAPACIDAD_FIREWALL.get(modelo, datos.get("capacidad", "-"))

    mrr_12 = TARIFA_SERVICIO_FIREWALL[modelo][12]
    mrr_24 = TARIFA_SERVICIO_FIREWALL[modelo][24]
    mrr_36 = TARIFA_SERVICIO_FIREWALL[modelo][36]

    while len(table.rows) < 3:
        table.add_row()

    eliminar_filas_extra(table, 3)

    if len(table.rows[0].cells) >= 1:
        escribir_celda(table.rows[0].cells[0], "PRECIARIO NGFW SOPHOS")

    headers = ["NGFW", "MRR 1 AÑO", "MRR 2 AÑOS", "MRR 3 AÑOS", "USUARIOS"]
    values = [
        modelo_tabla,
        formato_cop(mrr_12),
        formato_cop(mrr_24),
        formato_cop(mrr_36),
        str(capacidad),
    ]

    for idx, header in enumerate(headers):
        if idx < len(table.rows[1].cells):
            escribir_celda(table.rows[1].cells[idx], header)

    for idx, value in enumerate(values):
        if idx < len(table.rows[2].cells):
            escribir_celda(table.rows[2].cells[idx], value)

    return True


def construir_reemplazos_servicio(datos):
    fecha_actual = datetime.now().strftime("%d/%m/%Y")
    modelo = datos["modelo"]
    capacidad = datos.get("capacidad", CAPACIDAD_FIREWALL.get(modelo, "-"))

    modelo_con_capacidad = f"SOPHOS {modelo} ({capacidad} Usuarios aproximadamente)"

    reemplazos = {
        "NOMBRE CLIENTE": "Cliente",
        "XXXXXXX": "Cliente",
        "El cliente XXXXXXX requiere": "El cliente requiere",
        "Ciudad, fecha": f"Ciudad, {fecha_actual}",
        "Persona a quien va dirigida": "XXXX",
        "Cargo de la persona a quien va dirigida": "XXXX",
        "Ciudad en donde se oferta el servicio": "XXXX",
        "xxxxxxxxxxxx": "XXXX",
        "3xxxxxxxxx": "XXXX",
        "xxxx.xxxx@mc.net.co": "XXXX",
        "XXXXXX": "XXXX",
        "XXXXXXXXXXXXXX": "XXXX",
        "XXXXXXXX": "XXXX",
        "XXXXXXXX@mc.net.co": "XXXX",
        "Oferta venta de licenciamiento de seguridad informática NGFW": "Oferta de servicio de seguridad informática NGFW",
        "OFERTA venta de licenciamiento de seguridad informática NGFW": "Oferta de servicio de seguridad informática NGFW",
        "SOPHOS XGS 88 (20 Usuarios aproximadamente)": modelo_con_capacidad,
        "SOPHOS XGS88 (20 Usuarios aproximadamente)": modelo_con_capacidad,
        "Licenciamiento Xtream Protection": "Licenciamiento Xstream Protection",
        "Todos los equipos incluyen licencia Xtream Protection": "El equipo ofertado incluye licencia Xstream Protection",
    }

    return reemplazos


def generar_oferta_servicio_ngfw_word(datos):
    if not TEMPLATE_SERVICIO_NGFW_FILE.exists():
        return None

    doc = Document(str(TEMPLATE_SERVICIO_NGFW_FILE))
    reemplazos = construir_reemplazos_servicio(datos)

    reemplazar_texto_documento(doc, reemplazos)
    actualizar_tabla_preciario_servicio(doc, datos)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    imagen_firewall = obtener_imagen_firewall(datos["modelo"])

    buffer = reemplazar_imagen_y_xml_docx(
        docx_buffer=buffer,
        image_path=imagen_firewall,
        reemplazos_xml=reemplazos,
        media_target=WORD_MEDIA_FIREWALL_IMAGE_SERVICIO,
        image_format="JPEG",
    )

    buffer.seek(0)
    return buffer


# ============================================================
# WORD - ENDPOINT / SERVER / EDR / XDR
# ============================================================

def construir_reemplazos_endpoint():
    fecha_actual = datetime.now().strftime("%d/%m/%Y")

    return {
        "FUNDACION AMANECER": "Cliente",
        "FUNDACIÓN AMANECER": "Cliente",
        "INDEX TECNOLOGIA": "Cliente",
        "INDEX TECNOLOGÍA": "Cliente",
        "NOMBRE CLIENTE": "Cliente",
        "XXXXXXX": "Cliente",
        "El cliente FUNDACION AMANECER requiere": "El cliente requiere",
        "El cliente FUNDACIÓN AMANECER requiere": "El cliente requiere",
        "Pereira, 21 de nov. de 25": f"Pereira, {fecha_actual}",
        "Ciudad, fecha": f"Ciudad, {fecha_actual}",
        "Edinson Rodriguez": "XXXX",
        "Edinson Rodríguez": "XXXX",
        "Edinson.rodriguez@mc.net.co": "XXXX",
        "Edinson.Rodriguez@mc.net.co": "XXXX",
        "edinson.rodriguez@mc.net.co": "XXXX",
        "xxxxxxxxxxxx": "XXXX",
        "3xxxxxxxxx": "XXXX",
        "xxxx.xxxx@mc.net.co": "XXXX",
        "XXXXXXXX@mc.net.co": "XXXX",
        "XXXXXXXXXXXXXX": "XXXX",
        "XXXXXXXX": "XXXX",
    }


def normalizar_titulo_endpoint(texto):
    texto = texto.replace("–", "-").replace("—", "-")
    return normalizar_texto(texto)


def obtener_titulos_endpoint():
    return {
        "endpoint_user": [
            "Sophos Endpoint – User",
            "Sophos Endpoint - User",
            "Endpoint – User",
            "Endpoint - User",
        ],
        "endpoint_server": [
            "Sophos Endpoint – Server",
            "Sophos Endpoint - Server",
            "Endpoint – Server",
            "Endpoint - Server",
        ],
        "edr_user": [
            "Sophos EDR – User",
            "Sophos EDR - User",
            "EDR – User",
            "EDR - User",
        ],
        "edr_server": [
            "Sophos EDR – Server",
            "Sophos EDR - Server",
            "EDR – Server",
            "EDR - Server",
        ],
        "xdr_user": [
            "Sophos XDR – User",
            "Sophos XDR - User",
            "XDR – User",
            "XDR - User",
        ],
        "xdr_server": [
            "Sophos XDR – Server",
            "Sophos XDR - Server",
            "XDR – Server",
            "XDR - Server",
        ],
    }


def detectar_seccion_endpoint(texto):
    texto_norm = normalizar_titulo_endpoint(texto)
    titulos = obtener_titulos_endpoint()

    for key, variantes in titulos.items():
        for variante in variantes:
            if normalizar_titulo_endpoint(variante) in texto_norm:
                return key

    return None


def es_fin_bloques_endpoint(texto):
    texto_norm = normalizar_texto(texto)
    patrones_fin = [
        "plataformassoportadas",
        "instalacion",
        "soportetecnico",
        "administracion",
        "ofertaeconomica",
        "propuestacomercial",
        "condicionescomerciales",
    ]
    return any(p in texto_norm for p in patrones_fin)


def obtener_texto_elemento(element):
    textos = []
    for node in element.iter():
        if node.tag.endswith("}t") and node.text:
            textos.append(node.text)
    return " ".join(textos).strip()


def eliminar_elemento(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def limpiar_secciones_endpoint(doc, secciones_a_conservar):
    body = doc.element.body
    elementos = list(body)

    seccion_actual = None
    eliminando = False

    for element in elementos:
        texto = obtener_texto_elemento(element)

        if not texto:
            if eliminando:
                eliminar_elemento(element)
            continue

        if es_fin_bloques_endpoint(texto):
            seccion_actual = None
            eliminando = False
            continue

        seccion_detectada = detectar_seccion_endpoint(texto)

        if seccion_detectada:
            seccion_actual = seccion_detectada
            eliminando = seccion_detectada not in secciones_a_conservar

            if eliminando:
                eliminar_elemento(element)

            continue

        if seccion_actual is not None and eliminando:
            eliminar_elemento(element)


def encontrar_tabla_economica_endpoint(doc):
    for table in doc.tables:
        texto = " ".join(cell.text.strip() for row in table.rows for cell in row.cells)
        texto_norm = normalizar_texto(texto)

        if "ofertaeconomica" in texto_norm:
            return table

        if "valorunitario" in texto_norm and "valortotal" in texto_norm:
            return table

        if "referencia" in texto_norm and "valortotal" in texto_norm:
            return table

        if "item" in texto_norm and "cantidad" in texto_norm and "vigencia" in texto_norm:
            return table

    return None


def actualizar_tabla_economica_endpoint(doc, componentes, venta_instalacion):
    table = encontrar_tabla_economica_endpoint(doc)

    if table is None:
        return False

    asegurar_columnas_tabla(table, 4)

    filas = []

    for componente in componentes:
        filas.append(
            [
                componente["descripcion"],
                str(componente["cantidad"]),
                f"{componente['vigencia']} meses",
                formato_cop(componente["venta_cop"]),
            ]
        )

    if venta_instalacion and venta_instalacion > 0:
        filas.append(
            [
                "Servicios profesionales de instalación y puesta en marcha",
                "1",
                "Única vez",
                formato_cop(venta_instalacion),
            ]
        )

    total = sum(c["venta_cop"] for c in componentes) + (venta_instalacion or 0)

    filas.append(
        [
            "Total a presentar al cliente",
            "-",
            "-",
            formato_cop(total),
        ]
    )

    filas_necesarias = 1 + len(filas)

    while len(table.rows) < filas_necesarias:
        table.add_row()

    eliminar_filas_extra(table, filas_necesarias)

    headers = ["Ítem", "Cantidad", "Vigencia", "Valor Total COP"]

    for idx, header in enumerate(headers):
        escribir_celda(table.rows[0].cells[idx], header)

    for row_idx, fila in enumerate(filas, start=1):
        for col_idx, valor in enumerate(fila):
            escribir_celda(table.rows[row_idx].cells[col_idx], valor)

    return True


def key_seccion_endpoint(familia_excel):
    mapping = {
        "Endpoint": "endpoint_user",
        "Server": "endpoint_server",
        "EDR User": "edr_user",
        "EDR Server": "edr_server",
        "XDR User": "xdr_user",
        "XDR Server": "xdr_server",
    }
    return mapping.get(familia_excel)


def generar_oferta_endpoint_word(datos):
    if not TEMPLATE_ENDPOINT_FILE.exists():
        return None

    doc = Document(str(TEMPLATE_ENDPOINT_FILE))

    reemplazos = construir_reemplazos_endpoint()

    reemplazar_texto_documento(doc, reemplazos)
    limpiar_secciones_endpoint(doc, datos["secciones"])
    actualizar_tabla_economica_endpoint(
        doc=doc,
        componentes=datos["componentes"],
        venta_instalacion=datos["venta_instalacion"],
    )

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    buffer = reemplazar_imagen_y_xml_docx(
        docx_buffer=buffer,
        image_path=None,
        reemplazos_xml=reemplazos,
        media_target=None,
    )

    buffer.seek(0)
    return buffer


# ============================================================
# CARGA LISTA DE PRECIOS
# ============================================================

@st.cache_data
def cargar_lista_precios():
    if not PRICE_FILE.exists():
        return None, None

    xls = pd.ExcelFile(PRICE_FILE)

    if "Lista_Precios" in xls.sheet_names:
        hoja = "Lista_Precios"
    else:
        hoja = xls.sheet_names[0]

    df = pd.read_excel(PRICE_FILE, sheet_name=hoja)
    df = df.dropna(how="all")

    return df, hoja


def mapear_columnas(df):
    columnas = {
        "familia": obtener_columna(df, ["Familia"]),
        "producto": obtener_columna(df, ["Producto", "Modelo", "Modelo / Producto", "Nombre Producto"]),
        "min": obtener_columna(df, ["Min Usuarios", "Min", "Rango Min", "Desde", "Cantidad Min"]),
        "max": obtener_columna(df, ["Max Usuarios", "Max", "Rango Max", "Hasta", "Cantidad Max"]),
        "vigencia": obtener_columna(df, ["Vigencia Meses", "Vigencia", "Meses"]),
        "msrp": obtener_columna(df, ["MSRP USD", "MSRP", "Precio MSRP USD"]),
        "dr": obtener_columna(df, ["Con DR USD", "DR USD", "Precio DR USD", "Con Descuento USD"]),
        "sku": obtener_columna(df, ["SKU", "Código", "Codigo"]),
    }

    obligatorias = ["familia", "producto", "vigencia", "msrp", "dr"]
    faltantes = [nombre for nombre in obligatorias if columnas[nombre] is None]

    return columnas, faltantes


# ============================================================
# BÚSQUEDA DE PRECIOS
# ============================================================

def buscar_precio_firewall(df, columnas, modelo_sugerido, vigencia):
    familia_col = columnas["familia"]
    producto_col = columnas["producto"]
    vigencia_col = columnas["vigencia"]
    msrp_col = columnas["msrp"]
    dr_col = columnas["dr"]
    sku_col = columnas["sku"]

    df_temp = df.copy()
    df_temp["_familia_norm"] = df_temp[familia_col].apply(normalizar_texto)
    df_temp["_producto_norm"] = df_temp[producto_col].apply(normalizar_texto)
    df_temp["_vigencia_num"] = df_temp[vigencia_col].apply(convertir_numero)

    modelo_norm = normalizar_texto(modelo_sugerido)

    filtro_equipo = (
        (df_temp["_familia_norm"] == normalizar_texto("Firewall"))
        & (df_temp["_producto_norm"].str.contains(modelo_norm, na=False))
        & (df_temp["_vigencia_num"] == 0)
    )

    filtro_licencia = (
        (df_temp["_familia_norm"] == normalizar_texto("Firewall"))
        & (df_temp["_producto_norm"].str.contains(modelo_norm, na=False))
        & (df_temp["_vigencia_num"] == vigencia)
    )

    fila_equipo = df_temp[filtro_equipo]
    fila_licencia = df_temp[filtro_licencia]

    if fila_equipo.empty and fila_licencia.empty:
        return None

    resultado = pd.concat([fila_equipo, fila_licencia], ignore_index=True)

    msrp_equipo_usd = fila_equipo[msrp_col].apply(convertir_numero).sum() if not fila_equipo.empty else 0
    dr_equipo_usd = fila_equipo[dr_col].apply(convertir_numero).sum() if not fila_equipo.empty else 0

    msrp_licencia_usd = fila_licencia[msrp_col].apply(convertir_numero).sum() if not fila_licencia.empty else 0
    dr_licencia_usd = fila_licencia[dr_col].apply(convertir_numero).sum() if not fila_licencia.empty else 0

    if dr_equipo_usd <= 0:
        dr_equipo_usd = msrp_equipo_usd

    if dr_licencia_usd <= 0:
        dr_licencia_usd = msrp_licencia_usd

    msrp_total = msrp_equipo_usd + msrp_licencia_usd
    dr_total = dr_equipo_usd + dr_licencia_usd

    skus = []
    if sku_col:
        skus = resultado[sku_col].dropna().astype(str).tolist()

    detalle_firewall = pd.DataFrame(
        [
            {
                "Concepto": "Equipo / Hardware",
                "Modelo": modelo_sugerido,
                "Vigencia": "Pago único",
                "MSRP USD": formato_usd(msrp_equipo_usd),
                "DR USD": formato_usd(dr_equipo_usd),
            },
            {
                "Concepto": f"Licencia {vigencia} meses",
                "Modelo": modelo_sugerido,
                "Vigencia": f"{vigencia} meses",
                "MSRP USD": formato_usd(msrp_licencia_usd),
                "DR USD": formato_usd(dr_licencia_usd),
            },
            {
                "Concepto": "Total equipo + licencia",
                "Modelo": modelo_sugerido,
                "Vigencia": f"{vigencia} meses",
                "MSRP USD": formato_usd(msrp_total),
                "DR USD": formato_usd(dr_total),
            },
        ]
    )

    return {
        "producto": modelo_sugerido,
        "msrp_usd": msrp_total,
        "dr_usd": dr_total,
        "msrp_unitario_usd": msrp_total,
        "dr_unitario_usd": dr_total,
        "msrp_equipo_usd": msrp_equipo_usd,
        "dr_equipo_usd": dr_equipo_usd,
        "msrp_licencia_usd": msrp_licencia_usd,
        "dr_licencia_usd": dr_licencia_usd,
        "skus": skus,
        "filas": resultado,
        "detalle_firewall": detalle_firewall,
    }


def buscar_precio_licencia(df, columnas, familia, cantidad, vigencia):
    familia_col = columnas["familia"]
    producto_col = columnas["producto"]
    vigencia_col = columnas["vigencia"]
    msrp_col = columnas["msrp"]
    dr_col = columnas["dr"]
    min_col = columnas["min"]
    max_col = columnas["max"]
    sku_col = columnas["sku"]

    df_temp = df.copy()
    df_temp["_familia_norm"] = df_temp[familia_col].apply(normalizar_texto)
    df_temp["_vigencia_num"] = df_temp[vigencia_col].apply(convertir_numero)

    resultado = df_temp[
        (df_temp["_familia_norm"] == normalizar_texto(familia))
        & (df_temp["_vigencia_num"] == vigencia)
    ]

    if min_col and max_col:
        resultado = resultado[
            (resultado[min_col].apply(convertir_numero) <= cantidad)
            & (resultado[max_col].apply(convertir_numero) >= cantidad)
        ]

    if resultado.empty:
        return None

    fila = resultado.iloc[0]

    msrp_unitario = convertir_numero(fila[msrp_col])
    dr_unitario = convertir_numero(fila[dr_col])

    if dr_unitario <= 0:
        dr_unitario = msrp_unitario

    skus = []
    if sku_col and not pd.isna(fila[sku_col]):
        skus = [str(fila[sku_col])]

    return {
        "producto": fila[producto_col],
        "familia": familia,
        "cantidad": cantidad,
        "vigencia": vigencia,
        "msrp_usd": msrp_unitario * cantidad,
        "dr_usd": dr_unitario * cantidad,
        "msrp_unitario_usd": msrp_unitario,
        "dr_unitario_usd": dr_unitario,
        "skus": skus,
        "filas": resultado,
    }


# ============================================================
# CÁLCULOS FINANCIEROS
# ============================================================

def calcular_financiero(msrp_usd, dr_usd, escenario, trm, tipo_instalacion):
    msrp_usd = convertir_numero(msrp_usd)
    dr_usd = convertir_numero(dr_usd)

    costo_base_usd = obtener_costo_usd_por_escenario(msrp_usd, dr_usd, escenario)

    if costo_base_usd <= 0:
        costo_base_usd = msrp_usd

    costo_licencias_cop = costo_base_usd * trm
    precio_venta_licencias_cop = calcular_precio_venta_desde_costo(costo_licencias_cop)

    if tipo_instalacion == "Instalación completa":
        costo_instalacion_cop = precio_venta_licencias_cop * PORCENTAJE_COSTO_INSTALACION
        precio_venta_instalacion_cop = calcular_precio_venta_desde_costo(costo_instalacion_cop)
    else:
        costo_instalacion_cop = 0
        precio_venta_instalacion_cop = 0

    precio_venta_total_cop = precio_venta_licencias_cop + precio_venta_instalacion_cop
    costo_total_cop = costo_licencias_cop + costo_instalacion_cop

    utilidad_bruta_cop = precio_venta_total_cop - costo_total_cop
    impuestos_cop = precio_venta_total_cop * IMPUESTOS
    gastos_cop = precio_venta_total_cop * GASTOS
    margen_contribucional_cop = utilidad_bruta_cop - impuestos_cop - gastos_cop
    carga_prestacional_cop = precio_venta_total_cop * CARGA_PRESTACIONAL_COMISION
    comision_cop = precio_venta_total_cop * COMISION
    utilidad_ebit_cop = margen_contribucional_cop - carga_prestacional_cop - comision_cop

    porcentaje_costo = costo_total_cop / precio_venta_total_cop if precio_venta_total_cop else 0
    porcentaje_utilidad_bruta = utilidad_bruta_cop / precio_venta_total_cop if precio_venta_total_cop else 0
    porcentaje_margen_contribucional = margen_contribucional_cop / precio_venta_total_cop if precio_venta_total_cop else 0
    porcentaje_ebit = utilidad_ebit_cop / precio_venta_total_cop if precio_venta_total_cop else 0
    ebit_sobre_costo = utilidad_ebit_cop / costo_total_cop if costo_total_cop else 0

    return {
        "costo_licencias_cop": costo_licencias_cop,
        "precio_venta_licencias_cop": precio_venta_licencias_cop,
        "costo_instalacion_cop": costo_instalacion_cop,
        "precio_venta_instalacion_cop": precio_venta_instalacion_cop,
        "precio_venta_total_cop": precio_venta_total_cop,
        "costo_total_cop": costo_total_cop,
        "utilidad_bruta_cop": utilidad_bruta_cop,
        "impuestos_cop": impuestos_cop,
        "gastos_cop": gastos_cop,
        "margen_contribucional_cop": margen_contribucional_cop,
        "carga_prestacional_cop": carga_prestacional_cop,
        "comision_cop": comision_cop,
        "utilidad_ebit_cop": utilidad_ebit_cop,
        "porcentaje_costo": porcentaje_costo,
        "porcentaje_utilidad_bruta": porcentaje_utilidad_bruta,
        "porcentaje_margen_contribucional": porcentaje_margen_contribucional,
        "porcentaje_ebit": porcentaje_ebit,
        "ebit_sobre_costo": ebit_sobre_costo,
    }


def calcular_venta_componente(precio, escenario, trm):
    costo_usd = obtener_costo_usd_por_escenario(
        precio["msrp_usd"],
        precio["dr_usd"],
        escenario,
    )

    costo_cop = costo_usd * trm
    venta_cop = calcular_precio_venta_desde_costo(costo_cop)

    return {
        "costo_usd": costo_usd,
        "costo_cop": costo_cop,
        "venta_cop": venta_cop,
    }


def calcular_detalle_venta_firewall(precio, escenario, trm):
    costo_equipo_usd = obtener_costo_usd_por_escenario(
        precio["msrp_equipo_usd"],
        precio["dr_equipo_usd"],
        escenario,
    )

    costo_licencia_usd = obtener_costo_usd_por_escenario(
        precio["msrp_licencia_usd"],
        precio["dr_licencia_usd"],
        escenario,
    )

    costo_equipo_cop = costo_equipo_usd * trm
    costo_licencia_cop = costo_licencia_usd * trm

    venta_equipo_cop = calcular_precio_venta_desde_costo(costo_equipo_cop)
    venta_licencia_cop = calcular_precio_venta_desde_costo(costo_licencia_cop)

    return {
        "costo_equipo_usd": costo_equipo_usd,
        "costo_licencia_usd": costo_licencia_usd,
        "costo_equipo_cop": costo_equipo_cop,
        "costo_licencia_cop": costo_licencia_cop,
        "venta_equipo_cop": venta_equipo_cop,
        "venta_licencia_cop": venta_licencia_cop,
        "venta_equipo_licencia_cop": venta_equipo_cop + venta_licencia_cop,
    }


def calcular_servicio_firewall(modelo, vigencia, tipo_instalacion, cantidad_firewalls=1):
    mrc_unitario = TARIFA_SERVICIO_FIREWALL.get(modelo, {}).get(vigencia, 0)
    nrc_unitario = NRC_SERVICIO_FIREWALL.get(tipo_instalacion, 0)

    mrc_total = mrc_unitario * cantidad_firewalls
    nrc_total = nrc_unitario * cantidad_firewalls
    valor_contrato_mrc = mrc_total * vigencia
    valor_total_contrato = valor_contrato_mrc + nrc_total

    return {
        "modelo": modelo,
        "vigencia": vigencia,
        "cantidad_firewalls": cantidad_firewalls,
        "mrc_unitario": mrc_unitario,
        "nrc_unitario": nrc_unitario,
        "mrc_total": mrc_total,
        "nrc_total": nrc_total,
        "valor_contrato_mrc": valor_contrato_mrc,
        "valor_total_contrato": valor_total_contrato,
    }


# ============================================================
# HEADER
# ============================================================

st.markdown('<div class="mc-header">', unsafe_allow_html=True)
header_col1, header_col2 = st.columns([1, 3])

with header_col1:
    if LOGO_FILE.exists():
        st.image(str(LOGO_FILE), width=300)
    else:
        st.warning("Logo no encontrado en assets/logo_mc.png")

with header_col2:
    st.markdown(
        """
        <div class="mc-title">Cotizador Sophos SMB</div>
        <div class="mc-subtitle">
            Herramienta interna de Media Commerce para estimación rápida de servicios,
            licenciamiento y soluciones de seguridad Sophos SMB.
        </div>
        <div class="mc-badge-row">
            <div class="mc-badge">Seguridad Gestionada</div>
            <div class="mc-badge">Sophos SMB</div>
            <div class="mc-badge">Firewall · Endpoint · EDR · XDR</div>
            <div class="mc-badge">Modelo financiero controlado</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# CARGA EXCEL
# ============================================================

df_precios, hoja_usada = cargar_lista_precios()

if df_precios is None:
    st.error(
        "No se encontró la lista de precios. "
        "Valida que el archivo exista en: data/Lista de precios SOPHOS SMB.xlsx"
    )
    st.stop()

columnas, faltantes = mapear_columnas(df_precios)

if faltantes:
    st.error(
        "No se pudieron identificar estas columnas obligatorias en el Excel: "
        + ", ".join(faltantes)
    )
    st.write("Columnas encontradas en el archivo:")
    st.write(list(df_precios.columns))
    st.stop()

st.markdown(
    f"""
    <div class="mc-alert">
        Lista de precios cargada correctamente desde la hoja: <strong>{hoja_usada}</strong>
    </div>
    """,
    unsafe_allow_html=True,
)


# ============================================================
# FORMULARIO PRINCIPAL
# ============================================================

st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Parámetros de cotización</div>', unsafe_allow_html=True)

col1, col2, col3, col4 = st.columns(4)

with col1:
    familia = st.selectbox("Familia", FAMILIAS_VALIDAS)

with col2:
    if familia == "Firewall":
        cantidad = st.number_input(
            "Cantidad de usuarios",
            min_value=1,
            max_value=10000,
            value=20,
            step=1,
        )
    elif familia in ["Server", "EDR Server", "XDR Server"]:
        cantidad_usuarios_endpoint = st.number_input(
            "Cantidad de servidores",
            min_value=1,
            max_value=10000,
            value=1,
            step=1,
        )
    else:
        cantidad_usuarios_endpoint = st.number_input(
            "Cantidad de usuarios",
            min_value=1,
            max_value=10000,
            value=100,
            step=1,
        )

with col3:
    vigencia = st.selectbox("Vigencia", VIGENCIAS_VALIDAS, index=0)

with col4:
    trm = st.number_input("TRM", min_value=1.0, value=3800.0, step=10.0)

tipo_venta_firewall = "Venta one shot"
cantidad_firewalls = 1
incluye_server = False
cantidad_servidores = 0

if familia == "Firewall":
    col_fw_1, col_fw_2, col_fw_3 = st.columns(3)

    with col_fw_1:
        tipo_venta_firewall = st.selectbox(
            "Tipo de venta Firewall",
            TIPOS_VENTA_FIREWALL,
            index=0,
        )

    with col_fw_2:
        cantidad_firewalls = st.number_input(
            "Cantidad de firewalls",
            min_value=1,
            max_value=5,
            value=1,
            step=1,
        )

    with col_fw_3:
        st.info("Para servicio mensual se usa tarifa MRC/NRC del marco tarifario.")

else:
    incluye_server = False
    cantidad_servidores = 0

    if familia in ["Endpoint", "EDR User", "XDR User"]:
        col_ep_1, col_ep_2 = st.columns(2)

        with col_ep_1:
            incluye_server = st.checkbox(
                "Incluir licencias Server relacionadas",
                value=False,
                help="La vigencia de Server será obligatoriamente la misma seleccionada para usuarios.",
            )

        with col_ep_2:
            if incluye_server:
                cantidad_servidores = st.number_input(
                    "Cantidad de servidores",
                    min_value=1,
                    max_value=10000,
                    value=1,
                    step=1,
                )

    elif familia in ["Server", "EDR Server", "XDR Server"]:
        st.info("Cotización exclusiva de licencias Server. La vigencia seleccionada aplica a esta solución.")

col5, col6 = st.columns(2)

with col5:
    escenario = st.selectbox(
        "Estado de oportunidad",
        ["Sin DR aprobado", "Con DR aprobado"],
        index=1,
    )

with col6:
    if familia == "Firewall" and tipo_venta_firewall == "Servicio mensual":
        tipo_instalacion = st.selectbox(
            "Implementación",
            TIPOS_INSTALACION_FIREWALL_SERVICIO,
            index=0,
        )

    elif familia == "Firewall" and tipo_venta_firewall == "Venta one shot":
        tipo_instalacion = st.selectbox(
            "Tipo de instalación",
            TIPOS_INSTALACION_FIREWALL_VENTA,
            index=1,
        )

    else:
        tipo_instalacion = st.selectbox(
            "Tipo de instalación",
            TIPOS_INSTALACION_LICENCIAS,
            index=1,
        )

st.markdown("</div>", unsafe_allow_html=True)


# ============================================================
# FLUJO FIREWALL
# ============================================================

if familia == "Firewall":
    modelo_sugerido, capacidad_sugerida = sugerir_firewall(cantidad)

    if modelo_sugerido is None:
        st.error(
            "La cantidad ingresada supera el marco SMB para Firewall. "
            "Este caso requiere revisión técnica/comercial especial."
        )
        st.stop()

    st.markdown(
        f"""
        <div class="mc-alert">
            Equipo sugerido: <strong>Sophos {modelo_sugerido}</strong>
            para hasta <strong>{capacidad_sugerida} usuarios Full UTM</strong>.
        </div>
        """,
        unsafe_allow_html=True,
    )

    precio = buscar_precio_firewall(
        df=df_precios,
        columnas=columnas,
        modelo_sugerido=modelo_sugerido,
        vigencia=vigencia,
    )

    if precio is None:
        st.error("No se encontró precio para la combinación seleccionada.")
        st.stop()

    if tipo_venta_firewall == "Servicio mensual":
        servicio = calcular_servicio_firewall(
            modelo=modelo_sugerido,
            vigencia=vigencia,
            tipo_instalacion=tipo_instalacion,
            cantidad_firewalls=cantidad_firewalls,
        )

        st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
        st.markdown('<div class="mc-section-title">Resultado general · Servicio mensual</div>', unsafe_allow_html=True)

        s1, s2, s3, s4 = st.columns(4)

        with s1:
            render_result_card("MRC mensual", formato_cop(servicio["mrc_total"]), "Cargo mensual recurrente", "blue")

        with s2:
            render_result_card("NRC instalación", formato_cop(servicio["nrc_total"]), tipo_instalacion, "yellow")

        with s3:
            render_result_card(
                f"Valor MRC {vigencia} meses",
                formato_cop(servicio["valor_contrato_mrc"]),
                "Valor acumulado del servicio",
                "blue",
            )

        with s4:
            render_result_card("Total contrato", formato_cop(servicio["valor_total_contrato"]), "MRC contrato + NRC", "green")

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
        st.markdown('<div class="mc-section-title">Generar oferta comercial</div>', unsafe_allow_html=True)

        datos_oferta_servicio = {
            "modalidad": "Servicio mensual",
            "modelo": modelo_sugerido,
            "usuarios": cantidad,
            "capacidad": capacidad_sugerida,
            "vigencia": vigencia,
            "tipo_instalacion": tipo_instalacion,
            "cantidad_firewalls": cantidad_firewalls,
            "mrc_unitario": servicio["mrc_unitario"],
            "mrc_total": servicio["mrc_total"],
            "nrc_unitario": servicio["nrc_unitario"],
            "nrc_total": servicio["nrc_total"],
            "valor_contrato_mrc": servicio["valor_contrato_mrc"],
            "valor_total_contrato": servicio["valor_total_contrato"],
        }

        if not TEMPLATE_SERVICIO_NGFW_FILE.exists():
            st.error("No se encontró assets/oferta_servicio_ngfw.docx")
        else:
            archivo_servicio = generar_oferta_servicio_ngfw_word(datos_oferta_servicio)

            st.download_button(
                label="📄 Generar oferta de servicio NGFW en Word",
                data=archivo_servicio,
                file_name=nombre_archivo_oferta(modelo_sugerido, "Servicio mensual"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
        st.markdown('<div class="mc-section-title">Detalle del servicio firewall</div>', unsafe_allow_html=True)

        tabla_servicio = pd.DataFrame(
            [
                {"Concepto": "Equipo sugerido", "Valor": f"Sophos {modelo_sugerido}"},
                {"Concepto": "Cantidad de usuarios", "Valor": cantidad},
                {"Concepto": "Capacidad de referencia", "Valor": f"Hasta {capacidad_sugerida} usuarios Full UTM"},
                {"Concepto": "Cantidad de firewalls", "Valor": cantidad_firewalls},
                {"Concepto": "Vigencia", "Valor": f"{vigencia} meses"},
                {"Concepto": "Implementación", "Valor": tipo_instalacion},
                {"Concepto": "MRC unitario", "Valor": formato_cop(servicio["mrc_unitario"])},
                {"Concepto": "MRC total mensual", "Valor": formato_cop(servicio["mrc_total"])},
                {"Concepto": "NRC unitario", "Valor": formato_cop(servicio["nrc_unitario"])},
                {"Concepto": "NRC total", "Valor": formato_cop(servicio["nrc_total"])},
                {"Concepto": "Valor MRC contrato", "Valor": formato_cop(servicio["valor_contrato_mrc"])},
                {"Concepto": "Total contrato", "Valor": formato_cop(servicio["valor_total_contrato"])},
            ]
        )

        st.dataframe(tabla_servicio, use_container_width=True, hide_index=True)
        st.markdown("</div>", unsafe_allow_html=True)

        st.stop()

    # Venta one shot firewall
    msrp_usd = precio["msrp_usd"]
    dr_usd = precio["dr_usd"]

    resultado = calcular_financiero(
        msrp_usd=msrp_usd,
        dr_usd=dr_usd,
        escenario=escenario,
        trm=trm,
        tipo_instalacion=tipo_instalacion,
    )

    detalle_venta_firewall = calcular_detalle_venta_firewall(
        precio=precio,
        escenario=escenario,
        trm=trm,
    )

    st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
    st.markdown('<div class="mc-section-title">Resultado general</div>', unsafe_allow_html=True)

    m1, m2, m3, m4 = st.columns(4)

    with m1:
        render_result_card("Venta equipo", formato_cop(detalle_venta_firewall["venta_equipo_cop"]), "Hardware Sophos", "blue")

    with m2:
        render_result_card(
            f"Venta licencia {vigencia} meses",
            formato_cop(detalle_venta_firewall["venta_licencia_cop"]),
            "Suscripción seleccionada",
            "blue",
        )

    with m3:
        render_result_card("Venta instalación", formato_cop(resultado["precio_venta_instalacion_cop"]), tipo_instalacion, "yellow")

    with m4:
        render_result_card("Total cliente", formato_cop(resultado["precio_venta_total_cop"]), "Valor total estimado", "green")

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
    st.markdown('<div class="mc-section-title">Generar oferta comercial</div>', unsafe_allow_html=True)

    datos_oferta_utm = {
        "modalidad": "Venta one shot",
        "modelo": modelo_sugerido,
        "usuarios": cantidad,
        "capacidad": capacidad_sugerida,
        "vigencia": vigencia,
        "tipo_instalacion": tipo_instalacion,
        "venta_equipo": detalle_venta_firewall["venta_equipo_cop"],
        "venta_licencia": detalle_venta_firewall["venta_licencia_cop"],
        "venta_instalacion": resultado["precio_venta_instalacion_cop"],
        "total_cliente": resultado["precio_venta_total_cop"],
    }

    if not TEMPLATE_OFERTA_UTM_FILE.exists():
        st.error("No se encontró assets/oferta_venta_utm.docx")
    else:
        archivo_word = generar_oferta_venta_utm_word(datos_oferta_utm)

        st.download_button(
            label="📄 Generar oferta de venta UTM en Word",
            data=archivo_word,
            file_name=nombre_archivo_oferta(modelo_sugerido, "Venta one shot"),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    st.markdown("</div>", unsafe_allow_html=True)

    st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
    st.markdown('<div class="mc-section-title">Detalle Firewall: equipo y licencia</div>', unsafe_allow_html=True)

    detalle_fw = precio["detalle_firewall"].copy()

    detalle_fw["MSRP COP"] = [
        formato_cop(precio["msrp_equipo_usd"] * trm),
        formato_cop(precio["msrp_licencia_usd"] * trm),
        formato_cop(precio["msrp_usd"] * trm),
    ]

    detalle_fw["DR COP"] = [
        formato_cop(precio["dr_equipo_usd"] * trm),
        formato_cop(precio["dr_licencia_usd"] * trm),
        formato_cop(precio["dr_usd"] * trm),
    ]

    detalle_fw["Venta COP"] = [
        formato_cop(detalle_venta_firewall["venta_equipo_cop"]),
        formato_cop(detalle_venta_firewall["venta_licencia_cop"]),
        formato_cop(detalle_venta_firewall["venta_equipo_licencia_cop"]),
    ]

    st.dataframe(detalle_fw, use_container_width=True, hide_index=True)
    st.markdown("</div>", unsafe_allow_html=True)

    st.stop()


# ============================================================
# FLUJO ENDPOINT / SERVER / EDR / XDR
# ============================================================

componentes_a_cotizar = []

if familia == "Endpoint":
    componentes_a_cotizar.append(
        {
            "familia_excel": "Endpoint",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos Endpoint – User",
        }
    )

    if incluye_server:
        componentes_a_cotizar.append(
            {
                "familia_excel": "Server",
                "cantidad": cantidad_servidores,
                "descripcion": "Sophos Endpoint – Server",
            }
        )

elif familia == "Server":
    componentes_a_cotizar.append(
        {
            "familia_excel": "Server",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos Endpoint – Server",
        }
    )

elif familia == "EDR User":
    componentes_a_cotizar.append(
        {
            "familia_excel": "EDR User",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos EDR – User",
        }
    )

    if incluye_server:
        componentes_a_cotizar.append(
            {
                "familia_excel": "EDR Server",
                "cantidad": cantidad_servidores,
                "descripcion": "Sophos EDR – Server",
            }
        )

elif familia == "EDR Server":
    componentes_a_cotizar.append(
        {
            "familia_excel": "EDR Server",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos EDR – Server",
        }
    )

elif familia == "XDR User":
    componentes_a_cotizar.append(
        {
            "familia_excel": "XDR User",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos XDR – User",
        }
    )

    if incluye_server:
        componentes_a_cotizar.append(
            {
                "familia_excel": "XDR Server",
                "cantidad": cantidad_servidores,
                "descripcion": "Sophos XDR – Server",
            }
        )

elif familia == "XDR Server":
    componentes_a_cotizar.append(
        {
            "familia_excel": "XDR Server",
            "cantidad": cantidad_usuarios_endpoint,
            "descripcion": "Sophos XDR – Server",
        }
    )


precios_componentes = []

for componente in componentes_a_cotizar:
    precio_componente = buscar_precio_licencia(
        df=df_precios,
        columnas=columnas,
        familia=componente["familia_excel"],
        cantidad=componente["cantidad"],
        vigencia=vigencia,
    )

    if precio_componente is None:
        st.error(f"No se encontró precio para {componente['familia_excel']}.")
        st.stop()

    precio_componente["descripcion"] = componente["descripcion"]
    precios_componentes.append(precio_componente)


msrp_total = sum(p["msrp_usd"] for p in precios_componentes)
dr_total = sum(p["dr_usd"] for p in precios_componentes)

resultado = calcular_financiero(
    msrp_usd=msrp_total,
    dr_usd=dr_total,
    escenario=escenario,
    trm=trm,
    tipo_instalacion=tipo_instalacion,
)

detalle_componentes = []

for p in precios_componentes:
    venta = calcular_venta_componente(p, escenario, trm)

    detalle_componentes.append(
        {
            "familia": p["familia"],
            "descripcion": p["descripcion"],
            "producto": p["producto"],
            "cantidad": p["cantidad"],
            "vigencia": p["vigencia"],
            "msrp_usd": p["msrp_usd"],
            "dr_usd": p["dr_usd"],
            "venta_cop": venta["venta_cop"],
            "costo_cop": venta["costo_cop"],
        }
    )


st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Resultado general Endpoint</div>', unsafe_allow_html=True)

cols = st.columns(4)

with cols[0]:
    render_result_card(
        "Venta componente principal",
        formato_cop(detalle_componentes[0]["venta_cop"]),
        detalle_componentes[0]["descripcion"],
        "blue",
    )

with cols[1]:
    if len(detalle_componentes) > 1:
        render_result_card(
            "Venta componente adicional",
            formato_cop(detalle_componentes[1]["venta_cop"]),
            detalle_componentes[1]["descripcion"],
            "blue",
        )
    else:
        render_result_card(
            "Componente adicional",
            formato_cop(0),
            "No incluido",
            "blue",
        )

with cols[2]:
    render_result_card(
        "Venta instalación",
        formato_cop(resultado["precio_venta_instalacion_cop"]),
        tipo_instalacion,
        "yellow",
    )

with cols[3]:
    render_result_card(
        "Total cliente",
        formato_cop(resultado["precio_venta_total_cop"]),
        "Valor total estimado",
        "green",
    )

st.markdown("</div>", unsafe_allow_html=True)


st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Generar oferta Endpoint / EDR / XDR</div>', unsafe_allow_html=True)

secciones = []

for componente in detalle_componentes:
    key = key_seccion_endpoint(componente["familia"])
    if key:
        secciones.append(key)

datos_endpoint_word = {
    "familia": familia,
    "vigencia": vigencia,
    "secciones": secciones,
    "componentes": detalle_componentes,
    "venta_instalacion": resultado["precio_venta_instalacion_cop"],
    "total_cliente": resultado["precio_venta_total_cop"],
}

if not TEMPLATE_ENDPOINT_FILE.exists():
    st.error("No se encontró assets/oferta_endpoint.docx")
else:
    archivo_endpoint = generar_oferta_endpoint_word(datos_endpoint_word)

    st.download_button(
        label="📄 Generar oferta Endpoint en Word",
        data=archivo_endpoint,
        file_name=nombre_archivo_oferta(familia, "Endpoint"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.markdown("</div>", unsafe_allow_html=True)


st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Detalle de componentes cotizados</div>', unsafe_allow_html=True)

tabla_componentes = []

for c in detalle_componentes:
    tabla_componentes.append(
        {
            "Producto": c["descripcion"],
            "Cantidad": c["cantidad"],
            "Vigencia": f"{c['vigencia']} meses",
            "MSRP USD": formato_usd(c["msrp_usd"]),
            "DR USD": formato_usd(c["dr_usd"]),
            "Valor COP": formato_cop(c["venta_cop"]),
        }
    )

if resultado["precio_venta_instalacion_cop"] > 0:
    tabla_componentes.append(
        {
            "Producto": "Instalación y puesta en marcha",
            "Cantidad": 1,
            "Vigencia": "Única vez",
            "MSRP USD": "-",
            "DR USD": "-",
            "Valor COP": formato_cop(resultado["precio_venta_instalacion_cop"]),
        }
    )

tabla_componentes.append(
    {
        "Producto": "Total a presentar al cliente",
        "Cantidad": "-",
        "Vigencia": "-",
        "MSRP USD": formato_usd(msrp_total),
        "DR USD": formato_usd(dr_total),
        "Valor COP": formato_cop(resultado["precio_venta_total_cop"]),
    }
)

st.dataframe(pd.DataFrame(tabla_componentes), use_container_width=True, hide_index=True)
st.markdown("</div>", unsafe_allow_html=True)


st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Resultado financiero del negocio</div>', unsafe_allow_html=True)

tabla_financiera = pd.DataFrame(
    [
        {
            "Concepto": "Precio de venta total",
            "%": "100,0%",
            "Valor COP": formato_cop(resultado["precio_venta_total_cop"]),
        },
        {
            "Concepto": "Costo total",
            "%": formato_porcentaje(resultado["porcentaje_costo"]),
            "Valor COP": formato_cop(resultado["costo_total_cop"]),
        },
        {
            "Concepto": "Utilidad bruta",
            "%": formato_porcentaje(resultado["porcentaje_utilidad_bruta"]),
            "Valor COP": formato_cop(resultado["utilidad_bruta_cop"]),
        },
        {
            "Concepto": "Impuestos",
            "%": formato_porcentaje(IMPUESTOS),
            "Valor COP": formato_cop(resultado["impuestos_cop"]),
        },
        {
            "Concepto": "Gastos",
            "%": formato_porcentaje(GASTOS),
            "Valor COP": formato_cop(resultado["gastos_cop"]),
        },
        {
            "Concepto": "Margen contribucional",
            "%": formato_porcentaje(resultado["porcentaje_margen_contribucional"]),
            "Valor COP": formato_cop(resultado["margen_contribucional_cop"]),
        },
        {
            "Concepto": "Carga prestacional de la comisión",
            "%": formato_porcentaje(CARGA_PRESTACIONAL_COMISION),
            "Valor COP": formato_cop(resultado["carga_prestacional_cop"]),
        },
        {
            "Concepto": "Comisión",
            "%": formato_porcentaje(COMISION),
            "Valor COP": formato_cop(resultado["comision_cop"]),
        },
        {
            "Concepto": "Utilidad EBIT",
            "%": formato_porcentaje(resultado["porcentaje_ebit"]),
            "Valor COP": formato_cop(resultado["utilidad_ebit_cop"]),
        },
        {
            "Concepto": "EBIT / Costo",
            "%": formato_porcentaje(resultado["ebit_sobre_costo"]),
            "Valor COP": "-",
        },
    ]
)

st.dataframe(tabla_financiera, use_container_width=True, hide_index=True)
st.markdown("</div>", unsafe_allow_html=True)


st.markdown('<div class="mc-panel">', unsafe_allow_html=True)
st.markdown('<div class="mc-section-title">Validación comercial</div>', unsafe_allow_html=True)

ebit_costo = resultado["ebit_sobre_costo"]

if ebit_costo >= 0.15:
    st.success("Negocio viable: EBIT / Costo igual o superior al 15%.")
elif ebit_costo >= 0.08:
    st.warning("Negocio revisable: rentabilidad aceptable, pero requiere validación comercial.")
else:
    st.error("Negocio sensible: rentabilidad baja. Requiere revisión antes de presentar al cliente.")

st.markdown("</div>", unsafe_allow_html=True)


with st.expander("Validación técnica del cálculo"):
    st.write("Componentes usados:")
    st.write(detalle_componentes)

    st.write("Resultado financiero:")
    st.write(resultado)

    st.write("Filas usadas del Excel:")
    for p in precios_componentes:
        st.write(f"Familia: {p['familia']}")
        st.dataframe(p["filas"], use_container_width=True)

st.markdown(
    '<div class="mc-footer-note">Media Commerce · Herramienta interna de estimación comercial · Seguridad Sophos SMB</div>',
    unsafe_allow_html=True,
)